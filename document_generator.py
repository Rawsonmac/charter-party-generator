from docx import Document
from io import BytesIO
from datetime import datetime

def generate_document(template_name, custom_terms):
    doc = Document()
    doc.add_heading(f"{template_name} Charter Party Agreement", 0)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_heading("Terms and Conditions", level=1)
    
    for key, value in custom_terms.items():
        doc.add_paragraph(f"{key}: {value}", style='List Bullet')
    
    # Save document to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
